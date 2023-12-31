import sys
from PyQt5.QtWidgets import QApplication, QLabel, QMainWindow, QLineEdit, QPushButton, QTextEdit, QMessageBox, QMenuBar, QAction, QComboBox, QDateTimeEdit, QVBoxLayout, QWidget, QFormLayout, QFileDialog, QListWidget, QDateEdit, QTimeEdit
from PyQt5.QtCore import QDateTime, QDate, QTime
from docx import Document
from docx.shared import Inches
import os
import datetime
import string

class DocGenerator:
    """Class to handle the document generation logic."""

    def generate_document(self, template_path, meeting_details):
        try:
            document = Document(template_path)

            # First, handle the image insertion
            if 'logo' in meeting_details and meeting_details['logo']:
                self.insert_image(document, meeting_details['logo'], '[LOGO]')

            # Replace text placeholders in the document body
            for paragraph in document.paragraphs:
                for key, value in meeting_details.items():
                    if key != 'logo':  # Skip the logo key
                        placeholder = f'[{key.upper()}]'
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)

            # Replace placeholders in the footer
            for section in document.sections:
                footer = section.footer
                for paragraph in footer.paragraphs:
                    for key, value in meeting_details.items():
                        placeholder = f'[{key.upper()}]'
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)

            return document
        except Exception as e:
            print(f"Error generating document: {e}")
            return None

    def insert_image(self, document, image_path, placeholder):
        print(f"Attempting to insert image from {image_path}")

        if not os.path.exists(image_path):
            print(f"Image file not found at {image_path}")
            return

        image_inserted = False

        # Iterate through each section's header
        for section in document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                if placeholder in paragraph.text:
                    print(f"Found placeholder in header: {paragraph.text}")
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Inches(1))  # Adjust the width as needed
                    image_inserted = True
                    break

            if image_inserted:
                break

        if not image_inserted:
            print(f"Placeholder '{placeholder}' not found in the document headers.")

    def save_document(self, document, company_name):
        # Sanitize the company name to remove punctuation but retain spaces
        sanitized_company_name = self.remove_punctuation(company_name)

        # Generate a filename using the current date
        date_str = datetime.datetime.now().strftime("%Y-%m-%d")
        filename = f"{date_str} {sanitized_company_name} Board Meetings.docx"
        file_path = os.path.join("generated_docs", filename)

        # Ensure the directory exists
        if not os.path.exists("generated_docs"):
            os.makedirs("generated_docs")

        # Save the document
        document.save(file_path)
        return file_path

    def remove_punctuation(self, text):
        # Method to remove punctuation from text but retain spaces
        return text.translate(str.maketrans('', '', string.punctuation))

    def sanitize_company_name(self, company_name):
        # Original method to remove punctuation and spaces from the company name
        sanitized_name = company_name.translate(str.maketrans('', '', string.punctuation))
        return sanitized_name.replace(' ', '')


class DocGeneratorApp(QMainWindow):
    DEFAULT_LOGO_PATH = 'assets/logo.png'

    def __init__(self):
        super().__init__()
        self.doc_generator = DocGenerator()
        self.logo_path = self.DEFAULT_LOGO_PATH
        self.initUI()

    def initUI(self):
        self.setWindowTitle('docGenerator')
        self.setGeometry(100, 100, 700, 800)

        # Create central widget and layout
        central_widget = QWidget(self)
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)

        # Menu bar setup
        self.setupMenuBar()

        # Form layout for inputs
        form_layout = QFormLayout()
        main_layout.addLayout(form_layout)

        # Logo Selection
        self.logo_selection = QComboBox()
        self.logo_selection.addItems(['Logo', 'Upload Logo'])
        self.logo_selection.currentIndexChanged.connect(self.on_logo_selection_changed)
        form_layout.addRow('Select Logo:', self.logo_selection)

        self.logo_upload = QPushButton('Upload Logo', self)
        self.logo_upload.clicked.connect(self.upload_logo)
        self.logo_upload.hide()
        form_layout.addRow(self.logo_upload)

        # Company Selection Combo Box
        self.company_selection = QComboBox()
        self.company_selection.addItems(['Ryan Kyle', 'imnotrk', '6AMG', 'Aqua Fish Villa', 'TL Pet Shop', 'Other...'])
        self.company_selection.currentIndexChanged.connect(self.on_company_selection_changed)
        form_layout.addRow('Select Company:', self.company_selection)

        # Custom Company Input Field
        self.custom_company_input = QLineEdit()
        self.custom_company_input.setPlaceholderText("Enter custom company name")
        self.custom_company_input.hide()
        form_layout.addRow(self.custom_company_input)

        # Company Address Selection
        self.company_address_selection = QComboBox()
        self.company_address_selection.addItems([
            'Sample address 1, State US 12345',
            'Sample address 2, State US 12345',
            'Sample address 3, State US 12345',
            'Other...'
        ])
        form_layout.addRow('Company Address:', self.company_address_selection)

        self.custom_company_address_input = QLineEdit()
        self.custom_company_address_input.setPlaceholderText("Enter custom company address")
        self.custom_company_address_input.hide()
        form_layout.addRow(self.custom_company_address_input)

        # Meeting Type Selection
        self.meeting_type_selection = QComboBox()
        self.meeting_type_selection.addItems(['Annual Board Meeting', 'Special Board Meeting'])
        form_layout.addRow('Meeting Type:', self.meeting_type_selection)

        # # Attendees List
        # self.attendees_list = QListWidget()
        # self.attendees_list.addItem('Chairman of the Board')
        # form_layout.addRow('Attendees:', self.attendees_list)

        # self.new_attendee_input = QLineEdit()
        # self.new_attendee_input.setPlaceholderText("Enter attendee's name")
        # form_layout.addRow(self.new_attendee_input)

        # self.add_attendee_button = QPushButton('Add Attendee', self)
        # self.add_attendee_button.clicked.connect(self.add_attendee)
        # form_layout.addRow(self.add_attendee_button)

        # Chairman Name
        self.chairman_name_input = QLineEdit('Ryan Kyle Ocampo')
        form_layout.addRow('Chairman Name:', self.chairman_name_input)

        # Separate Date Input
        self.date_input = QDateEdit()
        self.date_input.setDate(QDate.currentDate())
        form_layout.addRow('Select Date:', self.date_input)

        # Modified Time Selection
        self.time_selection = QComboBox()
        self.time_selection.addItems(["10:00 AM", "10:30 AM", "11:00 AM", "12:00 PM", "Custom Time"])
        form_layout.addRow('Select Time:', self.time_selection)

        # Custom Time Input
        self.time_input = QTimeEdit()
        self.time_input.setTime(QTime.currentTime())
        self.time_input.setDisabled(True)
        form_layout.addRow('Custom Time:', self.time_input)

        self.time_selection.currentIndexChanged.connect(self.handleTimeSelection)

        # Meeting Address Selection
        self.meeting_address_selection = QComboBox()
        self.meeting_address_selection.addItems([
            'Sample address 1, State US 12345',
            'Sample address 2, State US 12345',
            'Sample address 3, State US 12345',
            'Other...'
        ])
        form_layout.addRow('Meeting Address:', self.meeting_address_selection)

        self.custom_meeting_address_input = QLineEdit()
        self.custom_meeting_address_input.setPlaceholderText("Enter custom meeting address")
        self.custom_meeting_address_input.hide()
        form_layout.addRow(self.custom_meeting_address_input)

        # Discussion, Resolutions, and Closing Remarks
        self.discussion_input = QTextEdit('Discuss business conditions and plans.')
        form_layout.addRow('Discussion:', self.discussion_input)

        self.resolutions_input = QTextEdit()
        form_layout.addRow('Resolutions:', self.resolutions_input)

        self.closing_remarks_input = QTextEdit('The topics were discussed and covered. The meeting was adjourned.')
        form_layout.addRow('Closing Remarks:', self.closing_remarks_input)

        # Generate document button
        self.btn_generate = QPushButton('Generate Document', self)
        self.btn_generate.clicked.connect(self.generate_document_from_input)
        main_layout.addWidget(self.btn_generate)

    def on_company_selection_changed(self, index):
        if self.company_selection.currentText() == "Other...":
            self.custom_company_input.show()
        else:
            self.custom_company_input.hide()
            self.custom_company_input.clear()

    def handleTimeSelection(self, index):
        if self.time_selection.currentText() == "Custom Time":
            self.time_input.setDisabled(False)
        else:
            self.time_input.setDisabled(True)
            selected_time = QTime.fromString(self.time_selection.currentText(), 'hh:mm AP')
            self.time_input.setTime(selected_time)

    def on_address_selection_changed(self, index):
        if self.company_address_selection.currentText() == "Other...":
            self.custom_company_address_input.show()
        else:
            self.custom_company_address_input.hide()
            self.custom_company_address_input.clear()

        if self.meeting_address_selection.currentText() == "Other...":
            self.custom_meeting_address_input.show()
        else:
            self.custom_meeting_address_input.hide()
            self.custom_meeting_address_input.clear()

    def setupMenuBar(self):
        menu_bar = QMenuBar(self)
        self.setMenuBar(menu_bar)

        file_menu = menu_bar.addMenu('&File')
        file_menu.addAction(QAction('&Open', self))
        file_menu.addAction(QAction('&New', self))
        file_menu.addAction(QAction('&Exit', self))

        menu_bar.addMenu('&Edit')
        menu_bar.addMenu('&View')
        menu_bar.addMenu('&Help')

    # def add_attendee(self):
    #     attendee_name = self.new_attendee_input.text().strip()
    #     if attendee_name:
    #         self.attendees_list.addItem(attendee_name)
    #         self.new_attendee_input.clear()

    def on_logo_selection_changed(self, index):
        if self.logo_selection.currentText() == "Upload Logo":
            self.logo_upload.show()
        else:
            self.logo_upload.hide()
            self.logo_path = self.DEFAULT_LOGO_PATH

    def upload_logo(self):
        file_name, _ = QFileDialog.getOpenFileName(self, "Select Logo", "", "Image Files (*.png *.jpg *.jpeg)")
        if file_name:
            self.logo_path = file_name

    def generate_document_from_input(self):
        # Use selected or custom company name
        company_name = self.company_selection.currentText()
        if company_name == "Other...":
            company_name = self.custom_company_input.text()

        # # Compile the list of attendees
        # attendees = [self.attendees_list.item(i).text() for i in range(self.attendees_list.count())]
        # attendees_str = ', '.join(attendees)

        # Use selected or custom company address
        company_address = self.company_address_selection.currentText()
        if company_address == "Other...":
            company_address = self.custom_company_address_input.text()

        # Use selected or custom meeting address
        meeting_address = self.meeting_address_selection.currentText()
        if meeting_address == "Other...":
            meeting_address = self.custom_meeting_address_input.text()

        # Extract the date and time
        meeting_date = self.date_input.date().toString("yyyy-MM-dd")
        if self.time_selection.currentText() == "Custom Time":
            meeting_time = self.time_input.time().toString("hh:mm a")
        else:
            meeting_time = self.time_selection.currentText()
        meeting_year = self.date_input.date().year()

        # Populate the meeting_details dictionary
        meeting_details = {
            'logo': self.logo_path,
            'company_name': company_name,
            'date': meeting_date,
            'time': meeting_time,
            'meeting_address': meeting_address,
            'chairman_name': self.chairman_name_input.text(),
            'year': str(meeting_year),
            'discussion_topics': self.discussion_input.toPlainText(),
            'resolutions': self.resolutions_input.toPlainText(),
            'closing_remarks': self.closing_remarks_input.toPlainText(),
            'closing_time': meeting_time,  
            'company_address': company_address,
        }

        # Generate the document
        document = self.doc_generator.generate_document("templates\Template.docx", meeting_details)
        if document is not None:
            file_path = self.doc_generator.save_document(document, company_name)
            QMessageBox.information(self, "Success", f"Document saved successfully at {file_path}")
        else:
            QMessageBox.critical(self, "Error", "Failed to generate the document.")

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = DocGeneratorApp()
    ex.show()
    sys.exit(app.exec_())